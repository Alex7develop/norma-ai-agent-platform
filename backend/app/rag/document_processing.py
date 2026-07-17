"""Safe document text extraction and deterministic MVP chunking."""

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentProcessingError(ValueError):
    """Raised when an uploaded document cannot be safely processed."""


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    """Normalized text extracted from a supported file."""

    text: str
    file_type: str


SUPPORTED_TYPES = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
}


def parse_document(
    *,
    filename: str,
    data: bytes,
    max_pages: int = 500,
    max_characters: int = 5_000_000,
) -> ParsedDocument:
    """Extract text without executing active document content."""

    suffix = Path(filename).suffix.lower()
    file_type = SUPPORTED_TYPES.get(suffix)
    if file_type is None:
        raise DocumentProcessingError(f"Unsupported document type: {suffix or 'none'}")
    if not data:
        raise DocumentProcessingError("Document is empty")

    try:
        if file_type == "pdf":
            reader = PdfReader(BytesIO(data))
            if reader.is_encrypted:
                raise DocumentProcessingError("Encrypted PDF files are not supported")
            if len(reader.pages) > max_pages:
                raise DocumentProcessingError("PDF page limit exceeded")
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        elif file_type == "docx":
            document = DocxDocument(BytesIO(data))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        else:
            text = data.decode("utf-8-sig")
    except DocumentProcessingError:
        raise
    except (ValueError, UnicodeDecodeError, OSError) as exc:
        raise DocumentProcessingError("Document content is invalid") from exc

    normalized = text.replace("\x00", "").strip()
    if not normalized:
        raise DocumentProcessingError("Document contains no extractable text")
    if len(normalized) > max_characters:
        raise DocumentProcessingError("Extracted text limit exceeded")
    return ParsedDocument(text=normalized, file_type=file_type)


def chunk_text(
    text: str,
    *,
    chunk_size: int,
    overlap: int,
) -> list[str]:
    """Split text deterministically while preferring natural boundaries."""

    if chunk_size <= 0 or overlap < 0 or overlap >= chunk_size:
        raise ValueError("Chunk size and overlap are invalid")

    normalized = re.sub(r"[ \t]+", " ", text)
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        proposed_end = min(start + chunk_size, len(normalized))
        end = proposed_end
        if proposed_end < len(normalized):
            boundary = max(
                normalized.rfind("\n", start, proposed_end),
                normalized.rfind(". ", start, proposed_end),
                normalized.rfind(" ", start, proposed_end),
            )
            if boundary > start + chunk_size // 2:
                end = boundary + 1

        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return chunks
